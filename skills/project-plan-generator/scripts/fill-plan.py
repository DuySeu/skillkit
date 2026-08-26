#!/usr/bin/env python3
"""Fill project-plan DOCX from approved Markdown using the bundled Word template."""

import argparse
import re
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def parse_frontmatter(text: str) -> Tuple[Dict[str, str], str]:
    """Parse optional YAML-style frontmatter and return body text."""
    if not text.startswith("---"):
        return {}, text
    end = text.find("---", 3)
    if end == -1:
        return {}, text
    frontmatter: Dict[str, str] = {}
    block = text[3:end].strip()
    for line in block.splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        frontmatter[key.strip()] = value.strip().strip('"').strip("'")
    return frontmatter, text[end + 3:].lstrip()


def parse_sections(body: str) -> Tuple[Dict[str, str], Dict[str, str]]:
    """Split markdown by headings; return section bodies and full heading titles."""
    sections: Dict[str, str] = {}
    titles: Dict[str, str] = {}
    current_key: Optional[str] = None
    current_lines: List[str] = []

    for line in body.splitlines():
        if line.startswith("## ") or line.startswith("### "):
            if current_key is not None:
                sections[current_key] = "\n".join(current_lines).strip()
            marker = line[3:] if line.startswith("## ") else line[4:]
            marker = marker.strip()
            token = marker.split()[0].rstrip(".")
            if re.match(r"^\d+(\.\d+)?$", token):
                current_key = token
                titles[current_key] = marker
            else:
                current_key = marker
                titles[current_key] = marker
            current_lines = []
        else:
            current_lines.append(line)

    if current_key is not None:
        sections[current_key] = "\n".join(current_lines).strip()
    return sections, titles


def parse_markdown_tables(block: str) -> List[List[List[str]]]:
    """Extract markdown pipe tables from a section block."""
    tables: List[List[List[str]]] = []
    lines = block.splitlines()
    index = 0
    while index < len(lines):
        if "|" not in lines[index]:
            index += 1
            continue
        if index + 1 >= len(lines) or "---" not in lines[index + 1]:
            index += 1
            continue
        table_lines: List[str] = []
        while index < len(lines) and "|" in lines[index]:
            table_lines.append(lines[index])
            index += 1
        rows: List[List[str]] = []
        for row_line in table_lines:
            if re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", row_line):
                continue
            cells = [cell.strip() for cell in row_line.strip().strip("|").split("|")]
            rows.append(cells)
        if rows:
            tables.append(rows)
    return tables


def strip_tables(block: str) -> str:
    """Return section prose with markdown tables removed."""
    lines = block.splitlines()
    output: List[str] = []
    index = 0
    while index < len(lines):
        if (
            "|" in lines[index]
            and index + 1 < len(lines)
            and "---" in lines[index + 1]
        ):
            while index < len(lines) and "|" in lines[index]:
                index += 1
            continue
        output.append(lines[index])
        index += 1
    return "\n".join(output).strip()


def heading_level(paragraph: Paragraph) -> Optional[int]:
    """Return heading level (1-9) or None if not a heading."""
    style_name = paragraph.style.name if paragraph.style else ""
    if not style_name.startswith("Heading"):
        return None
    suffix = style_name.replace("Heading ", "").strip()
    if suffix.isdigit():
        return int(suffix)
    return 1


def delete_paragraph(paragraph: Paragraph) -> None:
    """Remove a paragraph element from the document."""
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str,
    style_name: Optional[str] = None,
) -> Paragraph:
    """Insert a new paragraph immediately after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style_name:
        new_para.style = style_name
    return new_para


HEADING_ALIASES: Dict[str, str] = {
    "5.1": "Partner Technical Team",
    "5.2": "Resources & Cost Estimate",
}


def find_heading_paragraph(
    doc: Document,
    section_key: str,
    min_level: int = 1,
    title_hint: Optional[str] = None,
) -> Optional[Paragraph]:
    """Find a heading paragraph matching section key or alias text."""
    pattern = re.compile(rf"^\s*{re.escape(section_key)}[\s.]")
    alias = HEADING_ALIASES.get(section_key, "")
    hint = title_hint or alias

    for paragraph in doc.paragraphs:
        level = heading_level(paragraph)
        if level is None or level < min_level:
            continue
        text = paragraph.text.strip()
        if pattern.match(text):
            return paragraph
        if hint and hint.lower() in text.lower():
            return paragraph
    return None


def replace_heading_title(paragraph: Paragraph, new_title: str) -> None:
    """Replace entire heading paragraph text."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_title
    else:
        paragraph.add_run(new_title)


def has_child_sections(section_key: str, sections: Dict[str, str]) -> bool:
    """Return True if numbered child keys exist (e.g. 2.1 under 2)."""
    if not section_key.isdigit():
        return False
    prefix = f"{section_key}."
    return any(key.startswith(prefix) for key in sections if key != section_key)


def replace_section_body(
    doc: Document,
    section_key: str,
    body: str,
    min_level: int = 2,
    title_hint: Optional[str] = None,
) -> None:
    """Replace content under a heading until the next same-or-higher-level heading."""
    heading = find_heading_paragraph(
        doc, section_key, min_level=min_level, title_hint=title_hint
    )
    if heading is None:
        return

    heading_el = heading._element
    start_index = None
    for index, para in enumerate(doc.paragraphs):
        if para._element is heading_el:
            start_index = index
            break
    if start_index is None:
        return

    level = heading_level(doc.paragraphs[start_index]) or min_level
    end_index = len(doc.paragraphs)
    for index in range(start_index + 1, len(doc.paragraphs)):
        para = doc.paragraphs[index]
        para_level = heading_level(para)
        if para_level is not None and para_level <= level:
            end_index = index
            break

    for index in range(end_index - 1, start_index, -1):
        delete_paragraph(doc.paragraphs[index])

    prose = strip_tables(body)
    anchor = doc.paragraphs[start_index]
    if not prose:
        return

    for line in prose.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        style = None
        text = stripped
        if stripped.startswith("- ") or stripped.startswith("* "):
            style = "List Paragraph"
            text = stripped[2:].strip()
        anchor = insert_paragraph_after(anchor, text, style_name=style)


def fill_table(table, rows: List[List[str]], skip_header: bool = True) -> None:
    """Write rows into a Word table, preserving the header row."""
    data_rows = rows[1:] if skip_header and len(rows) > 1 else rows
    start_row = 1 if skip_header and len(table.rows) > 1 else 0

    for offset, row_data in enumerate(data_rows):
        row_index = start_row + offset
        if row_index >= len(table.rows):
            break
        row = table.rows[row_index]
        for col_index, cell_text in enumerate(row_data):
            if col_index >= len(row.cells):
                break
            row.cells[col_index].text = cell_text


def apply_cover(doc: Document, frontmatter: Dict[str, str]) -> None:
    """Set cover lines from frontmatter (paragraph indices 2–5 in bundled template)."""
    mapping = [
        ("program", 2),
        ("engagement_title", 4),
        ("use_case_oneline", 5),
    ]
    for key, index in mapping:
        value = frontmatter.get(key)
        if not value or index >= len(doc.paragraphs):
            continue
        paragraph = doc.paragraphs[index]
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = value
        else:
            paragraph.add_run(value)

    title = frontmatter.get("cover_title", "PROJECT PLAN")
    if len(doc.paragraphs) > 3:
        paragraph = doc.paragraphs[3]
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = title
        else:
            paragraph.add_run(title)


def to_doc_heading(md_title: str) -> str:
    """Convert markdown heading title to template double-space numbering."""
    match = re.match(r"^(\d+\.\d+)\s+(.*)$", md_title.strip())
    if match:
        return f"{match.group(1)}  {match.group(2)}"
    return md_title


def apply_section_headings(
    doc: Document,
    titles: Dict[str, str],
) -> None:
    """Update Heading2 text from markdown heading titles."""
    for key, full_title in titles.items():
        if not re.match(r"^\d+\.\d+$", key):
            continue
        heading = find_heading_paragraph(
            doc, key, min_level=2, title_hint=full_title
        )
        if heading is None:
            continue
        replace_heading_title(heading, to_doc_heading(full_title))


def fill_tables_from_sections(doc: Document, sections: Dict[str, str]) -> None:
    """Map parsed markdown tables to fixed template table indices."""
    if "1.2" in sections:
        tables = parse_markdown_tables(sections["1.2"])
        if tables and len(doc.tables) > 0:
            fill_table(doc.tables[0], tables[0])

    milestone_keys = sorted(
        key for key in sections if re.match(r"^4\.\d+$", key) and key != "4.4"
    )
    for offset, key in enumerate(milestone_keys[:3]):
        table_index = 1 + offset
        if table_index >= len(doc.tables):
            break
        tables = parse_markdown_tables(sections[key])
        if tables:
            fill_table(doc.tables[table_index], tables[0])

    cost_key = None
    for candidate in ("4.4", "4.5"):
        if candidate in sections:
            cost_key = candidate
            break
    if cost_key:
        tables = parse_markdown_tables(sections[cost_key])
        if len(tables) >= 1 and len(doc.tables) > 4:
            fill_table(doc.tables[4], tables[0])
        if len(tables) >= 2 and len(doc.tables) > 5:
            fill_table(doc.tables[5], tables[1])

    if "5" in sections:
        tables = parse_markdown_tables(sections["5"])
        if len(tables) >= 1 and len(doc.tables) > 6:
            fill_table(doc.tables[6], tables[0])
        if len(tables) >= 2 and len(doc.tables) > 7:
            fill_table(doc.tables[7], tables[1])


def fill_document(
    template_path: Path,
    output_path: Path,
    frontmatter: Dict[str, str],
    sections: Dict[str, str],
    titles: Dict[str, str],
) -> None:
    """Copy template and apply markdown content."""
    shutil.copy2(template_path, output_path)
    doc = Document(str(output_path))

    apply_cover(doc, frontmatter)
    apply_section_headings(doc, titles)

    subsection_keys = sorted(
        key for key in sections if re.match(r"^\d+\.\d+$", key)
    )
    for key in subsection_keys:
        replace_section_body(
            doc,
            key,
            sections[key],
            min_level=2,
            title_hint=titles.get(key),
        )

    top_level_keys = sorted(
        key for key in sections if re.match(r"^\d+$", key)
    )
    for key in top_level_keys:
        if has_child_sections(key, sections):
            if sections.get(key):
                replace_section_body(
                    doc,
                    key,
                    sections[key],
                    min_level=1,
                    title_hint=titles.get(key),
                )
            continue
        replace_section_body(
            doc,
            key,
            sections[key],
            min_level=1,
            title_hint=titles.get(key),
        )

    fill_tables_from_sections(doc, sections)
    doc.save(str(output_path))


def main() -> None:
    """CLI entry: markdown path in, DOCX path out."""
    skill_root = Path(__file__).resolve().parents[1]
    default_template = skill_root / "assets" / "project-plan-template.docx"

    parser = argparse.ArgumentParser(
        description="Export project-plan Markdown to styled DOCX."
    )
    parser.add_argument("markdown", type=Path, help="Approved plan markdown path")
    parser.add_argument(
        "--template",
        type=Path,
        default=default_template,
        help="DOCX template path",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="Output DOCX (default: same path as MD with .docx)",
    )
    args = parser.parse_args()

    if not args.markdown.is_file():
        raise SystemExit(f"Markdown not found: {args.markdown}")
    if not args.template.is_file():
        raise SystemExit(
            f"Template not found: {args.template}. "
            "Run scripts/prepare-template.py first."
        )

    output = args.output
    if output is None:
        output = args.markdown.with_suffix(".docx")

    text = args.markdown.read_text(encoding="utf-8")
    frontmatter, body = parse_frontmatter(text)
    sections, titles = parse_sections(body)
    fill_document(args.template, output, frontmatter, sections, titles)
    print(f"Written: {output}")


if __name__ == "__main__":
    main()
